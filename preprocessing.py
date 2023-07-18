import nltk

nltk.download('punkt')
nltk.download('wordnet')
nltk.download('stopwords')

import re
from docx import Document

raw_data = Document('Transcripts.docx')

def add_curly_braces(paragraph_text):
    """Add curly braces to three-digit numbers not surrounded by square brackets or curly braces."""
    numbers = re.findall(r'(?<![\[{])\b(\d{3})\b(?![\]}])', paragraph_text)
    for number in numbers:
        transformed_number = '{' + number + '}'
        paragraph_text = re.sub(r'\b' + number + r'\b', transformed_number, paragraph_text)
    return paragraph_text

def handle_square_brackets(paragraph_text):
    """Handle numbers inside square brackets with optional spaces and dashes."""
    matches = re.findall(r'\[([\d\s,-]+)\]', paragraph_text)
    for match in matches:
        numbers = []
        for num_range in re.split(r',\s*|\s+', match):
            num_range = num_range.strip()
            if '-' in num_range:
                start, end = num_range.split('-')
                numbers.extend(range(int(start), int(end) + 1))
            else:
                numbers.append(int(num_range))

        transformed = '[' + ']['.join(map(str, numbers)) + ']'
        paragraph_text = paragraph_text.replace('[' + match + ']', transformed)
    return paragraph_text

for each_paragraph in raw_data.paragraphs:
    each_paragraph.text = add_curly_braces(each_paragraph.text)
    each_paragraph.text = handle_square_brackets(each_paragraph.text)

def extract_info(paragraph_text):
    """Extract speaker and sentences from a paragraph text."""
    speaker_match = re.search(r'\{(\d{3})\}', paragraph_text)
    sentence_match = re.search(r': (.*)', paragraph_text)

    if speaker_match and sentence_match:
        speaker = speaker_match.group(1)
        sentence = sentence_match.group(1)
        return speaker, sentence
    else:
        return None, None


def update_data(data, speaker, sentence):
    """Update data dictionary with extracted speaker, sentence."""
    if speaker:
        if speaker in data:
            data[speaker]['statements'] = ''.join(data[speaker]['statements']) + sentence

        else:
            data[speaker] = {'statements': [statement]}

data = {}

for each_paragraph in raw_data.paragraphs:
    speaker, statement = extract_info(each_paragraph.text)
    update_data(data, speaker, statement)

del data['000']

import string
from nltk.tokenize import word_tokenize
from collections import defaultdict
from punctuator import Punctuator
import numpy as np

# Load the punctuator model
p = Punctuator('Demo-Europarl-EN.pcl')

def handle_disfluencies(text):
    # List of common disfluencies
    disfluencies = ['uh', 'um', 'like', 'you know', 'so', 'actually', 'basically', 'seriously', 'literally']
    
    # Tokenize the text
    words = word_tokenize(text)
    
    # Remove disfluencies
    words = [word for word in words if word not in disfluencies]
    
    return ' '.join(words)

def track_made_up_words(text, subject):
    # Tokenize the text
    words = word_tokenize(text)
    
    # List of common English words
    english_words = set(w.lower() for w in nltk.corpus.words.words())
    
    # Find made-up words
    made_up_words = [word for word in words if word not in english_words]
    
    return made_up_words

def remove_capitalization_and_punctuation(text):
    # Convert to lowercase
    text = text.lower()
    
    # Remove punctuation
    text = text.translate(str.maketrans('', '', string.punctuation))
    
    return text

def add_capitalization_and_punctuation(text):
    # Use the Punctuator to add capitalization and punctuation
    text = p.punctuate(text)
    
    return text

for subject, subject_data in data.items():

    text = subject_data['statememts']
    made_up_words = track_made_up_words(text, subject)
    text = handle_disfluencies(text)
    text = remove_capitalization_and_punctuation(text)
    text = add_capitalization_and_punctuation(text)

    subject_data['statements'] = text
    subject_data['made-up-words'] = made_up_words

# Save the output
with open('output.txt', 'w') as f:
    f.write(str(data))

